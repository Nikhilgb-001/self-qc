#!/usr/bin/env python3
"""
Self-QC Automation Script

This script automates the extraction of key fields from a Word document,
compares them against manual entries in an Excel checklist, and highlights
any mismatches.

Dependencies:
  pip install python-docx pandas openpyxl

Usage:
  python self_qc_automation.py \
    --docx path/to/document.docx \
    --excel path/to/checklist.xlsx \
    --output path/to/output_qc.xlsx
"""
import argparse
import re
from docx import Document
import pandas as pd
import openpyxl
from openpyxl.styles import PatternFill


def extract_fields_from_docx(doc_path, fields):
    """
    Extract values for each field from the Word document.
    Looks for patterns like "FieldName: Value" in paragraphs and tables.
    """
    doc = Document(doc_path)
    # Aggregate all text
    full_text = []
    for p in doc.paragraphs:
        full_text.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            full_text.append("\t".join(cells))
    text = "\n".join(full_text)
    field_values = {}
    for field in fields:
        pattern = re.compile(re.escape(field) + r"\s*[:\-]\s*(.+)")
        match = pattern.search(text)
        field_values[field] = match.group(1).strip() if match else None
    return field_values


def process_excel(excel_path, field_values, output_path):
    """
    Reads the Excel checklist, fills in auto-extracted values, compares
to manual entries, and writes a result Excel with mismatches highlighted.
    """
    df = pd.read_excel(excel_path, engine='openpyxl')
    # Expect columns: 'Field', 'Status', 'Manual Value'
    auto_vals = []
    matches = []
    for _, row in df.iterrows():
        field = row.get('Field')
        status = row.get('Status')
        manual = row.get('Manual Value')
        auto = field_values.get(field) if status else None
        auto_vals.append(auto)
        if status:
            if auto is None:
                matches.append(False)
            else:
                matches.append(str(auto).strip().lower() == str(manual).strip().lower())
        else:
            matches.append(None)
    df['Auto Extracted Value'] = auto_vals
    df['Match'] = matches

    # Write to Excel and highlight mismatches
    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='QC_Result')
        wb = writer.book
        ws = writer.sheets['QC_Result']
        red_fill = PatternFill(start_color='FFFFC7CE', end_color='FFFFC7CE', fill_type='solid')
        # Highlight rows where Match is False
        for idx, match in enumerate(matches, start=2):  # Excel rows start at 1 plus header
            if match is False:
                for col in range(1, len(df.columns) + 1):
                    ws.cell(row=idx, column=col).fill = red_fill
    print(f"QC results written to {output_path}")


def main():
    parser = argparse.ArgumentParser(description='Self-QC Automation')
    parser.add_argument('--docx', required=True, help='Path to the Word document (.docx)')
    parser.add_argument('--excel', required=True, help='Path to the Excel checklist file')
    parser.add_argument('--output', required=True, help='Path for the output QC Excel file')
    args = parser.parse_args()

    # Read input Excel to get the list of fields
    df_fields = pd.read_excel(args.excel, engine='openpyxl')
    fields = df_fields['Field'].dropna().astype(str).tolist()

    # Extract from Word
    field_values = extract_fields_from_docx(args.docx, fields)

    # Process and output
    process_excel(args.excel, field_values, args.output)


if __name__ == '__main__':
    main()
